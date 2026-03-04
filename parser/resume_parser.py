"""
resume_parser.py
Handles extraction of raw text from PDF, DOCX, and TXT resume files.
Also detects resume sections (Skills, Experience, Education, etc.)
"""

import re
import io
from pathlib import Path


def extract_text_from_pdf(file) -> str:
    """Extract text from a PDF file object or path."""
    try:
        import pdfplumber
        if isinstance(file, (str, Path)):
            with pdfplumber.open(file) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        else:
            with pdfplumber.open(file) as pdf:
                text = "\n".join(page.extract_text() or "" for page in pdf.pages)
        return text.strip()
    except ImportError:
        raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")
    except Exception as e:
        return f"Error reading PDF: {str(e)}"


def extract_text_from_docx(file) -> str:
    """Extract text from a DOCX file object or path."""
    try:
        import docx
        if isinstance(file, (str, Path)):
            doc = docx.Document(file)
        else:
            doc = docx.Document(io.BytesIO(file.read()))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n".join(paragraphs)
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"


def extract_text(file, file_type: str) -> str:
    """
    Unified text extractor.
    file_type: 'pdf', 'docx', or 'txt'
    """
    file_type = file_type.lower().strip(".")
    if file_type == "pdf":
        return extract_text_from_pdf(file)
    elif file_type in ("docx", "doc"):
        return extract_text_from_docx(file)
    elif file_type == "txt":
        if isinstance(file, (str, Path)):
            return Path(file).read_text(encoding="utf-8")
        return file.read().decode("utf-8")
    else:
        raise ValueError(f"Unsupported file type: {file_type}")


# ── Section Detection ──────────────────────────────────────────────────────────

SECTION_HEADERS = {
    "contact":     r"\b(contact|personal\s+info|profile)\b",
    "summary":     r"\b(summary|objective|about\s+me|professional\s+summary)\b",
    "skills":      r"\b(skills?|technical\s+skills?|core\s+competencies|expertise)\b",
    "experience":  r"\b(experience|work\s+history|employment|professional\s+experience)\b",
    "education":   r"\b(education|academic|qualifications?|degree)\b",
    "projects":    r"\b(projects?|portfolio|work\s+samples?)\b",
    "certifications": r"\b(certifications?|certificates?|licenses?|accreditations?)\b",
    "awards":      r"\b(awards?|honors?|achievements?|accomplishments?)\b",
    "languages":   r"\b(languages?)\b",
    "interests":   r"\b(interests?|hobbies|activities)\b",
}


def detect_sections(text: str) -> dict:
    """
    Parse resume text into labelled sections.
    Returns a dict: {section_name: section_text}
    """
    lines = text.split("\n")
    sections = {}
    current_section = "header"
    current_lines = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            current_lines.append(line)
            continue

        matched_section = None
        for section_name, pattern in SECTION_HEADERS.items():
            if re.search(pattern, stripped, re.IGNORECASE) and len(stripped) < 60:
                matched_section = section_name
                break

        if matched_section:
            sections[current_section] = "\n".join(current_lines).strip()
            current_section = matched_section
            current_lines = []
        else:
            current_lines.append(line)

    sections[current_section] = "\n".join(current_lines).strip()
    return {k: v for k, v in sections.items() if v}


def extract_contact_info(text: str) -> dict:
    """Extract email, phone, LinkedIn, GitHub from resume text."""
    email = re.findall(r"[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}", text)
    phone = re.findall(r"(\+?\d[\d\s\-().]{7,}\d)", text)
    linkedin = re.findall(r"linkedin\.com/in/[\w\-]+", text, re.IGNORECASE)
    github = re.findall(r"github\.com/[\w\-]+", text, re.IGNORECASE)

    return {
        "email": email[0] if email else None,
        "phone": phone[0].strip() if phone else None,
        "linkedin": linkedin[0] if linkedin else None,
        "github": github[0] if github else None,
    }